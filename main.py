import os
import json
import glob
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI
from fastapi import Request
from pydantic import BaseModel
from fastapi import Body, FastAPI
from fastapi.responses import JSONResponse
from fastapi.responses import FileResponse
from fastapi import FastAPI, File, UploadFile
from starlette.background import BackgroundTask
from fastapi.middleware.cors import CORSMiddleware

# Run with this: uvicorn main:app --reload

# Delete temp files in tmp


def deleteTemp():
    files = glob.glob(tmpDir + "*.docx")
    for f in files:
        os.remove(f)


dir = os.path.dirname(__file__)
modelosPath = dir + "/modelos/"
tmpDir = dir + "/tmp/"


def json_to_doc(json):
    temp_doc = Document()
    name = ""
    for i in json:
        match i["type"]:
            case "name":
                print("Nome do documento: " + i["t"])
                name = i["t"]
            case "p":
                # print("paragraph: " + i[1]["p"])
                print("paragraph|ajuda: " + str(i["ajuda"]))
                print("paragraph|removivel: " + str(i["removivel"]))
                print("paragraph|text:" + i["c"][0]["t"])
                temp_doc.add_paragraph(i["c"][0]["t"])
            case "i":
                print("image: " + i["i"])
                print("image|ajuda: " + i["ajuda"])
                print("image|removivel: " + str(i["removivel"]))
            case "h":
                print("heading: " + i["c"][0]["t"])
                print("heading|ajuda: " + i["ajuda"])
                print("heading|removivel: " + str(i["removivel"]))
                temp_doc.add_heading(i["c"][0]["t"], level=int(i["htype"]))
            case "t":
                print(f"table| [ rows: {i['rows']}|cols: {i['cols']} ]")
                print(f"table|len arrays: {len(i['cells'])}")
                print(f"table|arrays: {i['cells']}")
                table = temp_doc.add_table(rows=i["rows"], cols=i["cols"])

                for t in i["cells"]:
                    if t == i["cells"][0]:
                        counter = 0
                        for t in i["cells"][0]:
                            hdr_cell = table.rows[0].cells
                            hdr_cell[counter].text = t
                            counter = counter + 1
                    else:
                        cell = table.add_row().cells
                        for b in range(len(t)):
                            cell[b].text = t[b]
                temp_doc.add_paragraph("")

    temp_doc.save(f"{tmpDir}{name}.docx")
    return name


app = FastAPI()

origins = [
    "http://localhost/*",
    "http://localhost:5173",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# root should return all models
# Whoever consuming it should either treat it or smth


@app.get("/")
async def root():
    m = []
    for x in os.listdir(dir + "/modelos"):
        if x.endswith(".json"):
            fileJson = open(modelosPath + x)
            m.append(json.load(fileJson))
    return m


@app.get("/JSON/{modelo}")
async def mandarModeloJSON(modelo):
    path = os.path.join(dir + "/modelos/",
                        '{arquivo}.json'.format(arquivo=modelo))
    print(path)
    if os.path.exists(path):
        fileJSON = open(path)
        return json.load(fileJSON)
    else:
        return {"ERROR": "Arquivo não existe!"}


@app.post("/docx")
async def mandarModeloDocx(payload: Request):
    jd = json.loads(await payload.body())
    print(jd)
    name = json_to_doc(jd)
    # {"Receiced" : await payload.body()}#
    return FileResponse(f"{tmpDir}{name}.docx", background=BackgroundTask(deleteTemp))


@app.post("/docx/{modelo}")
async def mockDocx(modelo):
    fileData = open(os.path.join(dir + "/modelos/",
                    '{arquivo}.json'.format(arquivo=modelo)))
    jd = json.load(fileData)
    json_to_doc(jd)
    return FileResponse(f"{tmpDir}{modelo}.docx", background=BackgroundTask(deleteTemp))
