# importing the module
import json
from docx import Document
from docx.shared import Inches

#TODO: the match case in the test function is wrong, 

#JSON test
# creating the JSON data as a string
# data = '{"Name" : "Romy", "Gender" : "Female"}'
 
# print("Datatype before deserialization : "
#       + str(type(data)))
  
# # deserializing the data
# data = json.loads(data)

# print("Datatype after deserialization : "
#       + str(type(data)))

fileData = open("./test.json")

jd = json.load(fileData)

# print(jd)

# print("Datatype after deserialization : "
#        + str(type(jd)))

# p = jd["1"]["type"]

# print(p)

def test_json_to_doc(json):
      temp_doc = Document()
      name = "tmp"
      for i in json.items():
            #print(i)
            #print(i[1])
            #print(i[1]["type"])
            #match list(i[1].keys())[0]:
            match i[1]["type"]:
                  case "name":
                        print("Nome do documento: " + i[1]["t"])
                        name = i[1]["t"]
                  case "p":
                        #print("paragraph: " + i[1]["p"])
                        print("paragraph|ajuda: " + str(i[1]["ajuda"]))
                        print("paragraph|removivel: " + str(i[1]["removivel"]))
                        print("paragraph|text:" + i[1]["c"][0]["t"])
                        temp_doc.add_paragraph(i[1]["c"][0]["t"])
                  case "i":
                        print("image: " + i[1]["i"])
                        print("image|ajuda: " + i[1]["ajuda"])
                        print("image|removivel: " + str(i[1]["removivel"]))
                  case "h":
                        print("heading: " + i[1]["c"][0]["t"])
                        print("heading|ajuda: " + i[1]["ajuda"])
                        print("heading|removivel: " + str(i[1]["removivel"]))
                        temp_doc.add_heading(i[1]["c"][0]["t"], level=int(i[1]["htype"]))
                  case "t":
                        print(f"table| [ rows: {i[1]['rows']}|cols: {i[1]['cols']} ]")
                        print(f"table|len arrays: {len(i[1]['cells'])}")
                        print(f"table|arrays: {i[1]['cells']}")
                        table = temp_doc.add_table(rows=i[1]["rows"], cols=i[1]["cols"])

                        for t in i[1]["cells"]:
                              if t == i[1]["cells"][0]:
                                    counter = 0
                                    for t in i[1]["cells"][0]:
                                          hdr_cell = table.rows[0].cells
                                          hdr_cell[counter].text = t
                                          counter = counter + 1
                              else:
                                    cell = table.add_row().cells
                                    for b in range(len(t)):
                                          cell[b].text = t[b]

      temp_doc.save(f"{name}.docx")

test_json_to_doc(jd)